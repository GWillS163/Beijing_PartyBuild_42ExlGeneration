from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

correct_A = [
'非常好！配置详细，思路明确，现象完整。   	排版继续保持',
'一如既往的好, 注释圈点足以展示你的思维清晰',
'配置详细，现象基本完整。',
'写的不错,很好很好~'
]
correct_ASub = [
'截图详细 排版有待改进，可以尝试高亮，（用图粘配置，或者调整字体 能有效提升卷面）',
'现象基本完整，注释不足请改进。',
]
correct_Bsum = ['B+']
correct_C = ['C']
def edit_docx(filename,score):
    comment_scor = None
    if score == 'A':
        comment_scor = correct_A
    elif score == 'A-':
        comment_scor = correct_ASub
    elif score == 'B+':
        comment_scor = correct_Bsum
    elif score == 'C':
        comment_scor = correct_C
    else:
        print('分数值错误')
        return
    document = Document(filename)

    # document.add_paragraph('my paragraph') # 普通文本

    document.styles['Normal'].font.name = u'楷体'      #可换成word里面任意字体
    p = document.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER    #段落文字居中设置
    run = p.add_run(u'评语:'+ str(random.choice(comment_scor)))
    run.font.color.rgb = RGBColor(random.choices([i for i in range(255)])[0], random.choices([i for i in range(255)])[0], random.choices([i for i in range(255)])[0])             #颜色设置，这里是用RGB颜色
    run.font.size = Pt(random.choice([24, 36, 27, 32, 30]))                 #字体大小设置，和word里面的字号相对应

    document.add_picture('D:\python_scripts\\\qytang_script\\auto_correct_qyt_homework\qyt_' + score + '.png')  # 图片和python文件不在同一个文件夹下面的时候，要补全文件地址
    try:
        document.save('C:\\Users\\admin\Downloads\\' + filename.split('.')[0] + '-' + score + '.docx')
        print('\t\t作业批阅完毕')
    except PermissionError:
        print('⚠权限错误-请关闭同名文件！⚠')
    except:
        print('☢未知异常☢')
# edit_docx('Firewall ASA-day7-rsec-pangxt.docx','A-')