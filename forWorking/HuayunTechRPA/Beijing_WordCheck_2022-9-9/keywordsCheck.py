# import a module that can read docx & doc file
import docx
import os
import re
import csv
# Get current datetime as filename
import datetime


def findFileByRegex(folderPath, filesRegex):
    """根据正则表达式查找文件"""
    files = os.listdir(folderPath)
    return [file for file in files if re.match(filesRegex, file)]


def parseKeywords(keywordStr):
    """解析关键词"""
    keywordStr = keywordStr.replace("，", ",")
    res = []
    for keywordTuple in keywordStr.split(","):
        res.append(keywordTuple.split())
    return res


def saveResult2csv(outputFileName, resultList):
    with open(outputFileName, "w", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(resultList)


def singleFileCheck(filePath, wordsList):
    """进行单个文件的检查"""
    # content = docx2txt.process(filePath)
    content = readDocx2Txt(filePath)
    wordCheckList = [filePath]
    for keyword in wordsList:
        # 一组关键字有多个
        flag = False
        for kw in keyword:
            if kw in content:
                # wordCheckList.append(kw)  # 如果有一个关键字在文档中，就算通过
                wordCheckList.append("√" + kw)  # 如果有一个关键字在文档中，就算通过
                flag = True
                break
        # 如果一组关键字都不在文档中
        if not flag:
            wordCheckList.append("×关键词:"+".".join(keyword)+" 均不存在")  # 如果没有一个关键字在文档中，就算不通过
    return wordCheckList


def readDocx2Txt(param):
    """
    Read a docx file and return a string.
    :param param: a docx file path.
    :return: a string.
    """
    doc = docx.Document(param)
    # paragraphs
    text = ''
    for para in doc.paragraphs:
        text += para.text
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text += para.text
    return text


def main(filesPath, filesRegex, wordStr, outputPrefixName):
    """
    主程序
    :param filesPath:
    :param filesRegex:
    :param wordStr:
    :param outputPrefixName:
    :return:
    """
    files = findFileByRegex(filesPath, filesRegex)
    wordsList = parseKeywords(wordStr)
    os.chdir(filesPath)
    result = [["文件名", "关键词检查结果"]]
    for file in files:
        if "~$" in file:  # temp file
            continue
        # Method 1：使用全路径，输出文件在当前文件夹
        # filePath = os.path.join(filesPath, file)
        # result.append(singleFileCheck(filePath, wordsList))
        # Method 2: 使用相对路径，输出文件在Word文件夹
        result.append(singleFileCheck(file, wordsList))

    # 输出
    fileName = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    saveResult2csv(f"{outputPrefixName}_{fileName}.csv", result)
    print("Done!")
    return result

# need cancel the comment if you want to run this script
# res = main(filesPh, filesRe, kws, prefixName)
